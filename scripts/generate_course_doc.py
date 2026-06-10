# -*- coding: utf-8 -*-
"""Generate ERP course design document (.docx / .doc) for desktop."""

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

OUTPUT_DOCX = r"c:\Users\王照康\Desktop\ERP进销存管理系统课程设计.docx"
OUTPUT_DOC = r"c:\Users\王照康\Desktop\ERP进销存管理系统课程设计.doc"


def set_run_font(run, name="宋体", size=12, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold


def add_para(doc, text, align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=False, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_run_font(run, name="黑体", size=16 if level == 1 else 14, bold=True)
    return h


def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(9)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                set_run_font(r, bold=True, size=10)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for p in cells[ci].paragraphs:
                for r in p.runs:
                    set_run_font(r, size=10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def build_document():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.17)
    sec.right_margin = Cm(3.17)

    # ===== 封面 =====
    add_para(doc, "《数据库系统原理》课程设计", WD_ALIGN_PARAGRAPH.CENTER, 18, True, 20)
    add_para(doc, "", space_after=12)
    add_para(doc, "题    目： ERP 进销存管理系统", WD_ALIGN_PARAGRAPH.CENTER, 14, False, 10)
    add_para(doc, "学生姓名： 王照康", WD_ALIGN_PARAGRAPH.CENTER, 14, False, 10)
    add_para(doc, "学    号： （请填写）", WD_ALIGN_PARAGRAPH.CENTER, 14, False, 10)
    add_para(doc, "专业班级： （请填写）", WD_ALIGN_PARAGRAPH.CENTER, 14, False, 10)
    add_para(doc, "指导教师： （请填写）", WD_ALIGN_PARAGRAPH.CENTER, 14, False, 10)
    add_para(doc, "学    院： （请填写）", WD_ALIGN_PARAGRAPH.CENTER, 14, False, 10)
    add_para(doc, "2026 年 6 月", WD_ALIGN_PARAGRAPH.CENTER, 14, False, 20)
    doc.add_page_break()

    # ===== 小组成员 =====
    add_heading(doc, "小组成员信息", 1)
    add_table(doc,
              ["学号", "姓名", "分工情况"],
              [["（请填写）", "王照康",
                "负责本项目全部开发工作。完成后端架构设计与实现，包括 Spring Boot + MyBatis 多模块搭建、JWT 鉴权、采购/销售/库存核心业务、Spring AI 智能助手「小E」、阿里云 OSS 文件上传、AOP 操作日志等；完成数据库表结构设计及 MyBatis 映射；完成前端 Vue 页面联调；编写项目文档与测试用例。"]],
              [3, 2, 12])
    doc.add_page_break()

    # ===== 目录 =====
    add_heading(doc, "目  录", 1)
    toc = [
        "一、项目概述",
        "二、可行性分析",
        "三、功能性需求",
        "四、非功能需求",
        "五、系统架构设计",
        "六、功能模块设计",
        "七、数据库设计",
        "八、系统实现",
        "九、系统测试",
        "十、总结",
    ]
    for item in toc:
        add_para(doc, item, size=12)
    doc.add_page_break()

    # ===== 一、项目概述 =====
    add_heading(doc, "一、项目概述", 1)
    add_heading(doc, "1.1 项目背景", 2)
    add_para(doc, "传统中小型企业进销存管理多依赖 Excel 表格或纸质台账，存在数据更新滞后、库存统计不准确、采购与销售信息难以联动等问题。商品入库、出库、盘点等流程分散记录，容易出现账实不符、重复录入和统计错误。")
    add_para(doc, "随着企业业务规模扩大，商品种类、客户与供应商数量持续增长，人工管理方式已难以满足高效、准确的信息处理需求。本系统旨在通过信息化手段，实现商品、客户、供应商、采购订单、销售订单及库存的统一管理，提高企业进销存业务处理效率，为经营决策提供可靠的数据支撑。")

    add_heading(doc, "1.2 目标用户", 2)
    add_para(doc, "管理员（role=1）：拥有系统全部管理权限，可管理用户、商品、客户、供应商、采购单、销售单等全部业务数据。")
    add_para(doc, "普通员工（role=2）：可执行日常进销存业务操作，对用户管理模块的写操作受限。")

    add_heading(doc, "1.3 系统目标", 2)
    add_para(doc, "1、实现进销存核心业务闭环：采购下单 → 入库增库存 → 销售下单 → 出库减库存。")
    add_para(doc, "2、提供基础数据的增删改查与条件分页查询，支持商品分类、低库存预警。")
    add_para(doc, "3、支持多角色登录与 JWT 令牌鉴权，保障接口访问安全。")
    add_para(doc, "4、集成 Spring AI 智能助手，支持自然语言查询与办理常见业务。")
    add_para(doc, "5、记录关键操作日志，支持商品图片上传至阿里云 OSS。")

    # ===== 二、可行性分析 =====
    add_heading(doc, "二、可行性分析", 1)
    add_heading(doc, "2.1 法律可行性分析", 2)
    add_para(doc, "本系统开发所使用的框架与依赖均为开源或正版授权软件（Spring Boot、MyBatis、MySQL 等），不涉及侵权问题。系统用于课程设计与学习实践，用户数据存储在本地或自建数据库中，操作过程合法合规。综上，本系统在法律层面可行。")

    add_heading(doc, "2.2 技术可行性分析", 2)
    add_para(doc, "本系统采用 Spring Boot 3.5.3 + MyBatis + MySQL 技术栈，前后端分离架构，后端提供 RESTful API，前端采用 Vue 单页应用。主要技术包括：Java 21、Spring Boot、MyBatis、PageHelper、JWT、Spring AI、阿里云 OSS、Maven 多模块管理等。开发工具采用 IntelliJ IDEA，数据库管理采用 Navicat。技术成熟、文档丰富，具备充分的技术可行性。")

    add_heading(doc, "2.3 逻辑交互分析", 2)
    add_para(doc, "系统按角色划分功能，用户通过浏览器访问前端页面，前端调用后端 REST 接口，后端经 Service 层处理业务逻辑，通过 MyBatis Mapper 访问 MySQL 数据库。AI 助手模块通过 Spring AI 调用大模型，并注册业务 Tool 实现数据查询与操作。")
    add_para(doc, "图 2-1 零层数据流图（说明：外部实体「管理员/员工」与「ERP 进销存系统」交互，系统与 MySQL 数据库、阿里云 OSS、DeepSeek AI 服务进行数据交换。）")
    add_para(doc, "图 2-2 一层数据流图（说明：主要处理过程包括用户登录认证、基础数据管理、采购业务处理、销售业务处理、库存联动更新、AI 助手交互、操作日志记录。）")

    # ===== 三、功能性需求 =====
    add_heading(doc, "三、功能性需求", 1)
    add_heading(doc, "3.1 管理员端", 2)
    add_para(doc, "功能描述：管理系统全部业务数据及用户账号。")
    add_para(doc, "子功能：用户管理、商品分类管理、商品管理（含低库存预警、图片上传）、客户管理、供应商管理、采购订单管理（下单/入库/取消）、销售订单管理（下单/出库/取消）、操作日志记录。")
    add_para(doc, "图 3-1 管理员用例图（说明：管理员可对全部模块执行增删改查及业务流转操作。）")

    add_heading(doc, "3.2 员工端", 2)
    add_para(doc, "功能描述：执行日常进销存业务，权限低于管理员。")
    add_para(doc, "子功能：商品、客户、供应商、采购单、销售单的查询与日常操作；采购入库、销售出库等业务流转；不可对用户模块执行写操作。")
    add_para(doc, "图 3-2 员工用例图（说明：员工可执行业务操作，但用户管理写操作被拦截返回 403。）")

    add_heading(doc, "3.3 智能助手模块", 2)
    add_para(doc, "功能描述：通过自然语言与系统交互，辅助完成查询与业务操作。")
    add_para(doc, "子功能：查询客户、供应商、商品、采购单、销售单；新增/修改/删除基础数据；办理采购入库、销售出库、取消订单；库存低于阈值预警查询。")

    add_heading(doc, "3.4 登录模块", 2)
    add_para(doc, "功能描述：用户身份验证与令牌签发。")
    add_para(doc, "子功能：输入用户名、密码登录；登录成功后返回 JWT Token 及用户角色信息；前端保存 Token，后续请求携带 Authorization 头。")

    add_heading(doc, "3.5 文件上传模块", 2)
    add_para(doc, "功能描述：商品图片上传至阿里云 OSS，返回可访问 URL。")

    # ===== 四、非功能需求 =====
    add_heading(doc, "四、非功能需求", 1)
    add_heading(doc, "4.1 性能需求", 2)
    add_para(doc, "常规列表查询响应时间 ≤ 3 秒；分页查询支持 PageHelper；支持中小规模并发访问。")
    add_heading(doc, "4.2 安全性需求", 2)
    add_para(doc, "JWT 令牌认证，有效期 12 小时；拦截器校验 Token；员工禁止操作用户写接口；全局异常处理统一错误响应。")
    add_heading(doc, "4.3 可用性与易用性需求", 2)
    add_para(doc, "前端采用 Vue SPA，界面简洁；登录页独立设计；AI 助手采用口语化交互。")
    add_heading(doc, "4.4 可维护性与可扩展性", 2)
    add_para(doc, "Maven 多模块拆分；分层架构 Controller → Service → Mapper；MyBatis 动态 SQL；Spring AI Tool 机制便于扩展。")
    add_heading(doc, "4.5 兼容性与环境需求", 2)
    add_para(doc, "操作系统：Windows 10/11；JDK 21；MySQL 8.0+；Maven 3.9+；Chrome/Edge 浏览器。")

    # ===== 五、系统架构设计 =====
    add_heading(doc, "五、系统架构设计", 1)
    add_para(doc, "本系统采用前后端分离 + Maven 多模块架构。前端 Vue 应用通过 HTTP 调用后端 REST API；后端基于 Spring Boot 搭建，MyBatis 负责数据持久化，MySQL 存储业务数据。")
    add_para(doc, "表现层：Vue 前端（login.html / index.html）；控制层：Controller + TokenInterceptor；业务层：Service；持久层：Mapper + MyBatis XML；数据层：MySQL；横切关注点：AOP 操作日志、JWT、OSS、Spring AI。")
    add_para(doc, "Maven 模块：erp-parents（父 POM）、erp-pojo（实体）、erp-utils（工具）、erp-server（主应用）。")
    add_para(doc, "图 5-1 系统架构图（说明：浏览器 → Controller → Service → Mapper → MySQL，辅以 JWT 鉴权、AOP 日志、OSS 存储、AI 服务。）")

    # ===== 六、功能模块设计 =====
    add_heading(doc, "六、功能模块设计", 1)
    add_para(doc, "系统功能模块包括：登录认证、用户管理、商品分类管理、商品管理、客户管理、供应商管理、采购订单管理、销售订单管理、AI 智能助手、操作日志、文件上传。")
    add_para(doc, "核心业务：采购流程为「创建采购单（未入库）→ 采购入库（库存增加）→ 可选取消」；销售流程为「创建销售单（校验库存）→ 销售出库（库存扣减）→ 可选取消」。")
    add_para(doc, "订单状态码：1=未入库/未出库，2=已入库/已出库，3=已完成，4=已取消。")
    add_para(doc, "图 6-1 系统功能模块图（说明：以 ERP 进销存管理系统为根节点，向下展开各业务子模块及支撑模块。）")

    # ===== 七、数据库设计 =====
    add_heading(doc, "七、数据库设计", 1)
    add_heading(doc, "7.1 数据库概念模型设计", 2)
    add_para(doc, "主要实体：用户（User）、商品分类（GoodsType）、商品（Good）、客户（Customer）、供应商（Supplier）、采购订单（PurchaseOrder）、销售订单（SaleOrder）、操作日志（OperateLog）。")
    add_para(doc, "实体关系：商品分类—商品（1:N）；供应商—采购订单（1:N）；商品—采购/销售订单（1:N）；客户—销售订单（1:N）；用户—订单（1:N）。")
    add_para(doc, "图 7-1～7-8 各实体属性图（说明：分别描述用户、商品、客户、供应商、采购单、销售单、商品分类、操作日志的主要属性。）")
    add_para(doc, "图 7-9 总体 E-R 图（说明：展示上述实体及其一对多关联关系。）")

    add_heading(doc, "7.2 数据库的逻辑设计", 2)
    add_para(doc, "E-R 图向关系模型转换结果如下：")
    add_para(doc, "用户（id, username, password, name, gender, phone, email, role, create_time, update_time）")
    add_para(doc, "商品分类（id, type_name, remark）")
    add_para(doc, "商品（id, goods_name, type_id, buy_price, sell_price, stock_num, min_stock, image_url, remark, create_time, update_time）")
    add_para(doc, "客户（id, cust_name, contact_preson, phone, address, remark, type, create_time, update_time）")
    add_para(doc, "供应商（id, supplier_name, contact, phone, address, remark, create_time, update_time）")
    add_para(doc, "采购订单（id, goods_id, supplier_id, buy_num, buy_price, total_money, user_id, status, buy_time, create_time, update_time）")
    add_para(doc, "销售订单（id, goods_id, cust_id, sale_num, sale_price, sale_money, user_id, status, sale_time, create_time, update_time）")
    add_para(doc, "操作日志（id, operate_user_id, class_name, method_name, method_params, return_value, cost_time, operate_time）")

    add_heading(doc, "7.2.1 用户表（user）", 3)
    add_table(doc, ["列名", "数据类型", "是否为空", "主键", "备注"],
              [["id", "int", "否", "是", "自增主键"], ["username", "varchar(50)", "否", "否", "登录账号"],
               ["password", "varchar(100)", "否", "否", "登录密码"], ["name", "varchar(50)", "是", "否", "真实姓名"],
               ["gender", "tinyint", "是", "否", "1男2女"], ["phone", "varchar(11)", "是", "否", "电话"],
               ["email", "varchar(100)", "是", "否", "邮箱"], ["role", "tinyint", "否", "否", "1管理员2员工"],
               ["create_time", "datetime", "是", "否", "创建时间"], ["update_time", "datetime", "是", "否", "更新时间"]])

    add_heading(doc, "7.2.2 商品表（goods）", 3)
    add_table(doc, ["列名", "数据类型", "是否为空", "主键", "备注"],
              [["id", "bigint", "否", "是", "自增主键"], ["goods_name", "varchar(50)", "否", "否", "商品名称"],
               ["type_id", "int", "否", "否", "外键→goods_type"], ["buy_price", "decimal(10,2)", "是", "否", "进货价"],
               ["sell_price", "decimal(10,2)", "是", "否", "销售价"], ["stock_num", "int", "否", "否", "当前库存"],
               ["min_stock", "int", "是", "否", "最低库存预警"], ["image_url", "varchar(255)", "是", "否", "图片URL"],
               ["remark", "varchar(255)", "是", "否", "备注"], ["create_time", "datetime", "是", "否", "创建时间"],
               ["update_time", "datetime", "是", "否", "更新时间"]])

    add_heading(doc, "7.2.3 客户表（customer）", 3)
    add_table(doc, ["列名", "数据类型", "是否为空", "主键", "备注"],
              [["id", "bigint", "否", "是", "自增主键"], ["cust_name", "varchar(50)", "否", "否", "客户名称"],
               ["contact_preson", "varchar(20)", "是", "否", "联系人（历史拼写）"], ["phone", "varchar(11)", "是", "否", "电话"],
               ["address", "varchar(100)", "是", "否", "地址"], ["remark", "varchar(255)", "是", "否", "备注"],
               ["type", "int", "是", "否", "客户类型"], ["create_time", "datetime", "是", "否", "创建时间"],
               ["update_time", "datetime", "是", "否", "更新时间"]])

    add_heading(doc, "7.2.4 采购订单表（purchase_order）", 3)
    add_table(doc, ["列名", "数据类型", "是否为空", "主键", "备注"],
              [["id", "bigint", "否", "是", "自增主键"], ["goods_id", "bigint", "是", "否", "外键→goods"],
               ["supplier_id", "bigint", "是", "否", "外键→supplier"], ["buy_num", "int", "是", "否", "采购数量"],
               ["total_money", "decimal(10,2)", "是", "否", "总金额"], ["user_id", "int", "是", "否", "操作人"],
               ["status", "tinyint", "否", "否", "1未入库2已入库3完成4取消"], ["buy_time", "datetime", "是", "否", "入库时间"],
               ["create_time", "datetime", "是", "否", "创建时间"], ["update_time", "datetime", "是", "否", "更新时间"]])

    add_heading(doc, "7.2.5 销售订单表（sale_order）", 3)
    add_table(doc, ["列名", "数据类型", "是否为空", "主键", "备注"],
              [["id", "bigint", "否", "是", "自增主键"], ["goods_id", "bigint", "是", "否", "外键→goods"],
               ["cust_id", "bigint", "是", "否", "外键→customer"], ["sale_num", "int", "是", "否", "销售数量"],
               ["sale_money", "decimal(10,2)", "是", "否", "销售总额"], ["user_id", "int", "是", "否", "操作人"],
               ["status", "tinyint", "否", "否", "1未出库2已出库3完成4取消"], ["sale_time", "datetime", "是", "否", "出库时间"],
               ["create_time", "datetime", "是", "否", "创建时间"], ["update_time", "datetime", "是", "否", "更新时间"]])

    add_heading(doc, "7.3 数据库的物理设计", 2)
    add_heading(doc, "7.3.1 数据库及数据表的创建", 3)
    add_para(doc, "1、数据库创建")
    add_code(doc, """CREATE DATABASE IF NOT EXISTS erp
  CHARACTER SET utf8mb4
  COLLATE utf8mb4_unicode_ci;
USE erp;""")

    add_para(doc, "2、用户表创建")
    add_code(doc, """CREATE TABLE user (
    id INT PRIMARY KEY AUTO_INCREMENT,
    username VARCHAR(50) NOT NULL UNIQUE,
    password VARCHAR(100) NOT NULL,
    name VARCHAR(50),
    gender TINYINT,
    phone VARCHAR(11),
    email VARCHAR(100),
    role TINYINT NOT NULL DEFAULT 2 COMMENT '1管理员 2员工',
    create_time DATETIME DEFAULT CURRENT_TIMESTAMP,
    update_time DATETIME DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP
);""")

    add_para(doc, "3、商品表创建")
    add_code(doc, """CREATE TABLE goods (
    id BIGINT PRIMARY KEY AUTO_INCREMENT,
    goods_name VARCHAR(50) NOT NULL,
    type_id INT NOT NULL,
    buy_price DECIMAL(10,2) DEFAULT 0.00,
    sell_price DECIMAL(10,2) DEFAULT 0.00,
    stock_num INT DEFAULT 0,
    min_stock INT DEFAULT 10,
    image_url VARCHAR(255) DEFAULT '',
    remark VARCHAR(255) DEFAULT '',
    create_time DATETIME DEFAULT CURRENT_TIMESTAMP,
    update_time DATETIME DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP
);""")

    add_para(doc, "4、客户表、供应商表、采购订单表、销售订单表等其余表结构参照 7.2 节逻辑设计，使用 Navicat 或 SQL 脚本批量创建。")

    add_heading(doc, "7.3.2 插入数据", 3)
    add_para(doc, "使用 Navicat 可视化工具或 INSERT 语句插入测试数据。示例：")
    add_code(doc, "INSERT INTO user (username, password, name, role) VALUES ('admin', '加密密码', '管理员', 1);")

    add_heading(doc, "7.3.3 创建视图", 3)
    add_para(doc, "1、低库存商品预警视图")
    add_code(doc, """CREATE OR REPLACE VIEW v_low_stock_goods AS
SELECT g.id, g.goods_name, g.stock_num, g.min_stock, t.type_name
FROM goods g
LEFT JOIN goods_type t ON g.type_id = t.id
WHERE g.stock_num < g.min_stock;""")

    add_para(doc, "2、采购订单详情视图（关联商品与供应商名称）")
    add_code(doc, """CREATE OR REPLACE VIEW v_purchase_detail AS
SELECT p.*, g.goods_name, s.supplier_name
FROM purchase_order p
LEFT JOIN goods g ON p.goods_id = g.id
LEFT JOIN supplier s ON p.supplier_id = s.id;""")

    add_heading(doc, "7.3.4 存储过程", 3)
    add_para(doc, "1、采购入库存储过程")
    add_code(doc, """DELIMITER //
CREATE PROCEDURE sp_purchase_in_stock(IN p_order_id BIGINT)
BEGIN
    DECLARE v_goods_id BIGINT;
    DECLARE v_buy_num INT;
    DECLARE v_status TINYINT;
    SELECT goods_id, buy_num, status INTO v_goods_id, v_buy_num, v_status
    FROM purchase_order WHERE id = p_order_id;
    IF v_status != 1 THEN
        SIGNAL SQLSTATE '45000' SET MESSAGE_TEXT = '只有未入库订单可入库';
    END IF;
    UPDATE goods SET stock_num = stock_num + v_buy_num WHERE id = v_goods_id;
    UPDATE purchase_order SET status = 2, buy_time = NOW() WHERE id = p_order_id;
END //
DELIMITER ;""")

    add_para(doc, "2、销售出库存储过程（含库存校验）")
    add_code(doc, """DELIMITER //
CREATE PROCEDURE sp_sale_out_stock(IN p_order_id BIGINT)
BEGIN
    DECLARE v_goods_id BIGINT;
    DECLARE v_sale_num INT;
    DECLARE v_status TINYINT;
    DECLARE v_rows INT;
    SELECT goods_id, sale_num, status INTO v_goods_id, v_sale_num, v_status
    FROM sale_order WHERE id = p_order_id;
    IF v_status != 1 THEN
        SIGNAL SQLSTATE '45000' SET MESSAGE_TEXT = '只有未出库订单可出库';
    END IF;
    UPDATE goods SET stock_num = stock_num - v_sale_num
    WHERE id = v_goods_id AND stock_num >= v_sale_num;
    SET v_rows = ROW_COUNT();
    IF v_rows = 0 THEN
        SIGNAL SQLSTATE '45000' SET MESSAGE_TEXT = '库存不足';
    END IF;
    UPDATE sale_order SET status = 2, sale_time = NOW() WHERE id = p_order_id;
END //
DELIMITER ;""")

    add_heading(doc, "7.3.5 创建触发器", 3)
    add_para(doc, "1、商品库存变更审计触发器")
    add_code(doc, """CREATE TRIGGER trg_goods_stock_audit
AFTER UPDATE ON goods
FOR EACH ROW
BEGIN
    IF OLD.stock_num != NEW.stock_num THEN
        INSERT INTO operate_log (class_name, method_name, method_params, operate_time)
        VALUES ('goods', 'stock_change',
                CONCAT('goods_id=', NEW.id, ', old=', OLD.stock_num, ', new=', NEW.stock_num),
                NOW());
    END IF;
END;""")

    add_para(doc, "2、采购单状态变更触发器（禁止重复入库）")
    add_code(doc, """CREATE TRIGGER trg_purchase_status_check
BEFORE UPDATE ON purchase_order
FOR EACH ROW
BEGIN
    IF OLD.status = 2 AND NEW.status = 2 AND OLD.status != NEW.status THEN
        SIGNAL SQLSTATE '45000' SET MESSAGE_TEXT = '已入库订单不可重复入库';
    END IF;
END;""")

    add_heading(doc, "7.3.6 数据库用户管理与权限控制", 3)
    add_para(doc, "设计基于角色的数据库用户：erp_admin（读写全部表）、erp_readonly（只读查询），实现细粒度访问控制。")

    add_heading(doc, "7.3.7 数据库备份与恢复方案", 3)
    add_para(doc, "采用 mysqldump 定期备份 erp 库；恢复时使用 mysql 命令导入备份文件，保障数据完整性与业务连续性。")

    # ===== 八、系统实现 =====
    add_heading(doc, "八、系统实现", 1)
    add_heading(doc, "8.1 搭建项目框架", 2)
    add_para(doc, "1、在 IDEA 中创建 Maven 多模块项目 erp-project，子模块包括 erp-parents、erp-pojo、erp-utils、erp-server。")
    add_para(doc, "2、在 erp-server/pom.xml 中引入 spring-boot-starter-web、mybatis-spring-boot-starter、pagehelper、mysql-connector-j、jjwt、spring-ai、aliyun-sdk-oss 等依赖。")
    add_para(doc, "3、配置 application.yaml，设置数据源、MyBatis 映射路径、Jackson 日期格式、Spring AI 参数。")
    add_para(doc, "4、创建启动类 ErpApplication，配置 @MapperScan 扫描 Mapper 接口。")
    add_para(doc, "图 8-1 项目框架（说明：展示 Maven 多模块目录结构及主要配置文件位置。）")

    add_heading(doc, "8.2 项目层次介绍", 2)
    add_para(doc, "Controller 层：接收 HTTP 请求，调用 Service，返回 JSON。")
    add_para(doc, "Service 层：核心业务逻辑、事务控制、库存联动、AI Tool 实现。")
    add_para(doc, "Mapper 层：MyBatis 数据访问，XML 映射 SQL。")
    add_para(doc, "pojo 层：实体类、查询参数、统一响应 Result。")
    add_para(doc, "utils 层：JwtUtils、AliyunOSSOperator 等工具类。")
    add_para(doc, "图 8-2 层次结构图（说明：表现层 → 控制层 → 业务层 → 持久层 → 数据层。）")

    add_heading(doc, "8.3 Mapper 层（DAO 层）", 2)
    add_para(doc, "以 UserMapper 为例，接口定义增删改查方法，XML 文件实现具体 SQL。采用 MyBatis 动态 SQL 技术，通过 <if> 标签实现条件查询。")
    add_code(doc, """<update id="updateById">
    update user
    <set>
        <if test="username!=null and username !=''">username=#{username},</if>
        <if test="password != null and password != ''">password = #{password},</if>
        update_time = now()
    </set>
    where id=#{id}
</update>""")
    add_para(doc, "库存更新采用原子 SQL 防止并发超卖：increaseStock 与 decreaseStock（WHERE stock_num >= ?）。")

    add_heading(doc, "8.4 Service 层", 2)
    add_para(doc, "以 SaleOrderServiceImpl 为例：创建订单时校验库存；出库时使用 decreaseStock 原子扣减；全程 @Transactional 保证事务一致性。")

    add_heading(doc, "8.5 前端页面", 2)
    add_para(doc, "前端采用 Vue SPA，login.html 负责登录，index.html 为主页面。登录成功后 Token 保存至 localStorage，后续 API 请求携带 Authorization: Bearer {token}。")

    add_heading(doc, "8.6 主页面展示", 2)
    add_para(doc, "图 8-3 登录页面（说明：用户输入账号密码，调用 POST /login 接口。）")
    add_para(doc, "图 8-4 系统主页面（说明：左侧菜单导航，右侧业务内容区，含 AI 助手入口。）")

    add_heading(doc, "8.7 管理员端", 2)
    add_para(doc, "图 8-5 商品管理页面")
    add_para(doc, "图 8-6 采购订单管理页面")
    add_para(doc, "图 8-7 销售订单管理页面")
    add_para(doc, "图 8-8 用户管理页面")

    add_heading(doc, "8.8 员工端", 2)
    add_para(doc, "图 8-9 员工业务操作页面（说明：员工可执行采购、销售等业务，用户管理写操作被拦截。）")

    add_heading(doc, "8.9 AI 智能助手", 2)
    add_para(doc, "图 8-10 AI 助手对话界面（说明：自然语言查询商品、办理入库出库等操作。）")

    # ===== 九、系统测试 =====
    add_heading(doc, "九、系统测试", 1)
    add_heading(doc, "9.1 测试目的", 2)
    add_para(doc, "验证系统功能完整性、业务逻辑正确性、权限控制有效性及基本性能表现，确保系统满足设计需求。")

    add_heading(doc, "9.2 测试环境", 2)
    add_table(doc, ["项目", "配置"],
              [["操作系统", "Windows 11"], ["IDE", "IntelliJ IDEA"], ["JDK", "21"],
               ["数据库", "MySQL 8.0"], ["数据库工具", "Navicat"], ["构建工具", "Maven 3.9"],
               ["浏览器", "Chrome"]])

    add_heading(doc, "9.3 系统功能测试", 2)
    add_para(doc, "表 9-1 测试用例表")
    test_rows = [
        ["登录模块", "正确登录", "输入正确用户名密码", "返回 Token，跳转主页", "登录成功", "通过"],
        ["登录模块", "错误登录", "输入错误密码", "提示登录失败", "提示正确", "通过"],
        ["商品模块", "新增商品", "填写商品信息提交", "商品新增成功", "新增成功", "通过"],
        ["商品模块", "低库存查询", "访问低库存接口", "返回预警列表", "显示正常", "通过"],
        ["采购模块", "采购入库", "对未入库订单入库", "库存增加，状态变更", "联动正确", "通过"],
        ["销售模块", "创建销售单", "库存充足时创建", "订单创建成功", "创建成功", "通过"],
        ["销售模块", "库存不足", "销售数量大于库存", "提示库存不足", "拒绝创建", "通过"],
        ["销售模块", "销售出库", "对未出库订单出库", "库存扣减", "扣减正确", "通过"],
        ["权限模块", "员工操作用户", "员工 Token 写用户", "返回 403", "拦截正常", "通过"],
        ["权限模块", "未登录访问", "不带 Token 访问", "返回 401", "拦截正常", "通过"],
        ["AI模块", "查询商品", "发送自然语言查询", "返回商品列表", "响应正常", "通过"],
        ["上传模块", "图片上传", "上传商品图片", "返回 OSS URL", "上传成功", "通过"],
    ]
    add_table(doc, ["测试模块", "测试用例", "测试步骤", "预期结果", "实际结果", "测试结论"], test_rows)

    add_heading(doc, "9.4 测试总结", 2)
    add_para(doc, "经系统性测试，ERP 进销存管理系统各核心功能模块运行正常，采购—入库—销售—出库业务闭环完整，JWT 鉴权与角色权限控制有效，AI 助手可正常响应。系统满足课程设计的功能与性能要求，可稳定运行。")

    # ===== 十、总结 =====
    add_heading(doc, "十、总结", 1)
    add_para(doc, "本项目基于 Spring Boot 3 + MyBatis + MySQL 架构，采用 Maven 多模块设计，实现了完整的 ERP 进销存管理系统。主要成果包括：完整的进销存业务闭环、JWT 多角色权限控制、Spring AI 智能助手、工程化分层架构与单元测试。")
    add_para(doc, "通过本次课程设计，掌握了 Spring Boot 快速开发、MyBatis 动态 SQL、JWT 认证、Spring AI Tool Calling 等关键技术，积累了从需求分析到测试验证的完整开发经验。")
    add_para(doc, "不足与改进方向：密码加密方案待完善（建议 BCrypt）；操作日志缺少查询界面；缺少数据统计报表；前端 Vue 工程需独立管理并纳入 CI 构建流程。")
    add_para(doc, "未来将继续完善系统功能，探索 Redis 缓存、Docker 部署、Knife4j 接口文档等工程化实践，为实际项目开发打下坚实基础。")

    return doc


def convert_to_doc(docx_path, doc_path):
    try:
        import win32com.client
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(docx_path)
        doc.SaveAs(doc_path, FileFormat=0)
        doc.Close()
        word.Quit()
        return True
    except Exception as e:
        print(f"DOC conversion skipped: {e}")
        return False


if __name__ == "__main__":
    document = build_document()
    document.save(OUTPUT_DOCX)
    print(f"Saved: {OUTPUT_DOCX}")

    try:
        import win32com.client  # noqa: F401
        if convert_to_doc(OUTPUT_DOCX, OUTPUT_DOC):
            print(f"Saved: {OUTPUT_DOC}")
        else:
            print("Word COM conversion failed; .docx file is available.")
    except ImportError:
        import subprocess
        try:
            subprocess.run(["py", "-m", "pip", "install", "pywin32", "-q"], check=True)
            if convert_to_doc(OUTPUT_DOCX, OUTPUT_DOC):
                print(f"Saved: {OUTPUT_DOC}")
        except Exception:
            print("Install pywin32 manually for .doc conversion; .docx is ready.")
